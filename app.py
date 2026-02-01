# app.py - Flask Backend optimized for Render Free Tier
import os
import shutil
import time
from flask import Flask, render_template, request, jsonify, send_from_directory
from flask_cors import CORS
from dotenv import load_dotenv
from werkzeug.utils import secure_filename

# Load environment variables first
load_dotenv()

# Flask app setup
app = Flask(__name__, static_folder='static', template_folder='templates')
CORS(app)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'dev-secret-key')
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size

# Create necessary directories
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs('static', exist_ok=True)
os.makedirs('templates', exist_ok=True)

# Configuration
persist_directory = "chroma_db"
ALLOWED_EXTENSIONS = {'pdf', 'txt', 'docx'}

# Lazy import flags
HAS_LANGCHAIN = False
HAS_WATSONX = False

def initialize_dependencies():
    """Lazy initialization to save memory"""
    global HAS_LANGCHAIN, HAS_WATSONX
    
    try:
        # Try to import LangChain components
        from langchain_community.embeddings import HuggingFaceEmbeddings
        from langchain_community.vectorstores import Chroma
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        from langchain.chains import RetrievalQA
        from langchain.prompts import PromptTemplate
        from langchain.schema import Document
        HAS_LANGCHAIN = True
        print("✓ LangChain loaded")
    except ImportError as e:
        print(f"✗ LangChain not available: {e}")
    
    try:
        from langchain_ibm import WatsonxLLM
        HAS_WATSONX = True
        print("✓ WatsonxLLM loaded")
    except ImportError as e:
        print(f"✗ WatsonxLLM not available: {e}")

# Initialize dependencies
initialize_dependencies()

def get_embeddings():
    """Get embeddings model (lightweight)"""
    if not HAS_LANGCHAIN:
        return None
    try:
        # Use a very lightweight model
        from langchain_community.embeddings import HuggingFaceEmbeddings
        return HuggingFaceEmbeddings(
            model_name="all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
    except Exception as e:
        print(f"✗ Could not load embeddings: {e}")
        return None

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def check_ibm_credentials():
    """Verify IBM credentials"""
    IBM_API_KEY = os.getenv("IBM_API_KEY")
    IBM_PROJECT_ID = os.getenv("IBM_PROJECT_ID")
    IBM_URL = os.getenv("IBM_URL", "https://us-south.ml.cloud.ibm.com")
    
    if not IBM_API_KEY or not IBM_PROJECT_ID:
        return False, "IBM credentials not found in environment variables"
    return True, ""

def clear_database():
    """Clear vector database"""
    try:
        if os.path.exists(persist_directory):
            shutil.rmtree(persist_directory)
            print("Database cleared")
            return True
        return True  # Return True even if it doesn't exist
    except Exception as e:
        print(f"Error clearing database: {e}")
    return False

def extract_text_fallback(file_path, original_filename):
    """Simple text extraction for free tier"""
    documents = []
    
    if not HAS_LANGCHAIN:
        return documents
    
    from langchain.schema import Document
    
    try:
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return documents
            
        if original_filename.lower().endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                if content.strip():
                    doc = Document(
                        page_content=content,
                        metadata={"source": original_filename}
                    )
                    documents.append(doc)
                    
        elif original_filename.lower().endswith('.pdf'):
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                content = ""
                # Only read first 10 pages to save memory
                max_pages = min(10, len(reader.pages))
                for page_num in range(max_pages):
                    try:
                        page = reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            content += page_text + "\n\n"
                    except Exception as page_error:
                        print(f"Error reading page {page_num}: {page_error}")
                        continue
                
                if content.strip():
                    doc = Document(
                        page_content=content,
                        metadata={
                            "source": original_filename,
                            "pages_read": max_pages
                        }
                    )
                    documents.append(doc)
            except Exception as e:
                print(f"PDF extraction error: {e}")
                # Try alternative method
                try:
                    import fitz  # PyMuPDF
                    doc = fitz.open(file_path)
                    content = ""
                    max_pages = min(10, len(doc))
                    for page_num in range(max_pages):
                        page = doc.load_page(page_num)
                        content += page.get_text() + "\n\n"
                    doc.close()
                    
                    if content.strip():
                        doc = Document(
                            page_content=content,
                            metadata={
                                "source": original_filename,
                                "pages_read": max_pages
                            }
                        )
                        documents.append(doc)
                except Exception as e2:
                    print(f"Alternative PDF extraction also failed: {e2}")
                
        elif original_filename.lower().endswith('.docx'):
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
                if content.strip():
                    doc = Document(
                        page_content=content,
                        metadata={"source": original_filename}
                    )
                    documents.append(doc)
            except Exception as e:
                print(f"DOCX extraction error: {e}")
                
    except Exception as e:
        print(f"Extraction error: {e}")
    
    return documents

# Routes
@app.route('/')
def index():
    """Serve main page"""
    return render_template('webpage.html')

@app.route('/static/<path:filename>')
def serve_static(filename):
    """Serve static files"""
    return send_from_directory('static', filename)

@app.route('/healthz')
def health_check():
    """Health check endpoint"""
    return jsonify({
        "status": "healthy",
        "service": "document-qa",
        "memory": "ok"
    }), 200

@app.route('/api/status')
def status():
    """Check service status"""
    creds_ok, creds_msg = check_ibm_credentials()
    return jsonify({
        "status": "running",
        "langchain": HAS_LANGCHAIN,
        "watsonx": HAS_WATSONX,
        "database_exists": os.path.exists(persist_directory),
        "ibm_credentials": creds_ok,
        "max_file_size": "50MB",
        "supported_formats": list(ALLOWED_EXTENSIONS)
    })

@app.route('/api/clear-database', methods=['POST'])
def api_clear_database():
    """Clear the vector database"""
    try:
        success = clear_database()
        if success:
            return jsonify({
                'status': 'success',
                'message': 'Database cleared successfully'
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'Failed to clear database'
            }), 500
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@app.route('/api/upload-documents', methods=['POST'])
def upload_documents():
    """Handle document upload"""
    # Check dependencies
    if not HAS_LANGCHAIN:
        return jsonify({
            'status': 'error',
            'message': 'LangChain not available'
        }), 500
    
    # Check credentials
    creds_ok, creds_msg = check_ibm_credentials()
    if not creds_ok:
        return jsonify({
            'status': 'error',
            'message': creds_msg
        }), 400
    
    if 'files' not in request.files:
        return jsonify({
            'status': 'error',
            'message': 'No files uploaded'
        }), 400
    
    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        return jsonify({
            'status': 'error',
            'message': 'No files selected'
        }), 400
    
    # Clear old database to save memory
    clear_database()
    
    file_paths = []
    processed_files = []
    
    try:
        # Save uploaded files (limit to 2 files for free tier)
        for file in files[:2]:  # Limit to 2 files
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)
                file_paths.append((file_path, filename))
                processed_files.append(filename)
        
        if not file_paths:
            return jsonify({
                'status': 'error',
                'message': 'No valid files uploaded'
            }), 400
        
        print(f"Processing {len(file_paths)} files...")
        
        # Extract text
        all_docs = []
        for file_path, filename in file_paths:
            docs = extract_text_fallback(file_path, filename)
            all_docs.extend(docs)
            print(f"Extracted {len(docs)} documents from {filename}")
            
            # Clean up file immediately
            if os.path.exists(file_path):
                os.remove(file_path)
        
        if not all_docs:
            return jsonify({
                'status': 'error',
                'message': 'Could not extract text from documents'
            }), 500
        
        # Get chunk size from request
        chunk_size = request.form.get('chunk_size', 1000)
        chunk_overlap = request.form.get('chunk_overlap', 200)
        
        try:
            chunk_size = int(chunk_size)
            chunk_overlap = int(chunk_overlap)
        except:
            chunk_size = 1000
            chunk_overlap = 200
        
        # Split documents
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        splits = text_splitter.split_documents(all_docs)
        
        print(f"Created {len(splits)} chunks")
        
        # Create vector database
        embeddings = get_embeddings()
        if not embeddings:
            return jsonify({
                'status': 'error',
                'message': 'Could not initialize embeddings'
            }), 500
        
        from langchain_community.vectorstores import Chroma
        vectorstore = Chroma.from_documents(
            documents=splits,
            embedding=embeddings,
            persist_directory=persist_directory
        )
        vectorstore.persist()
        
        # Clear memory
        import gc
        gc.collect()
        
        return jsonify({
            'status': 'success',
            'message': f'Successfully processed {len(processed_files)} document(s)!',
            'data': {
                'files': processed_files,
                'chunks': len(splits),
                'documents': len(all_docs)
            }
        })
        
    except Exception as e:
        print(f"Upload error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@app.route('/api/ask-question', methods=['POST'])
def ask_question():
    """Handle questions"""
    if not HAS_LANGCHAIN:
        return jsonify({
            'status': 'error',
            'message': 'LangChain not available'
        }), 500
    
    if not HAS_WATSONX:
        return jsonify({
            'status': 'error',
            'message': 'Watsonx not available'
        }), 500
    
    creds_ok, creds_msg = check_ibm_credentials()
    if not creds_ok:
        return jsonify({
            'status': 'error',
            'message': creds_msg
        }), 400
    
    if not os.path.exists(persist_directory):
        return jsonify({
            'status': 'error',
            'message': 'Please upload documents first'
        }), 400
    
    data = request.get_json()
    if not data or 'question' not in data:
        return jsonify({
            'status': 'error',
            'message': 'No question provided'
        }), 400
    
    question = data['question'].strip()
    if not question:
        return jsonify({
            'status': 'error',
            'message': 'Please enter a question'
        }), 400
    
    try:
        # Load vectorstore
        embeddings = get_embeddings()
        if not embeddings:
            return jsonify({
                'status': 'error',
                'message': 'Could not load embeddings'
            }), 500
            
        from langchain_community.vectorstores import Chroma
        vectorstore = Chroma(
            persist_directory=persist_directory,
            embedding_function=embeddings
        )
        
        # Initialize LLM
        from langchain_ibm import WatsonxLLM
        llm = WatsonxLLM(
            model_id="ibm/granite-3-8b-instruct",
            url=os.getenv("IBM_URL", "https://us-south.ml.cloud.ibm.com"),
            apikey=os.getenv("IBM_API_KEY"),
            project_id=os.getenv("IBM_PROJECT_ID"),
            params={
                "temperature": 0.1,
                "max_new_tokens": 500,
                "min_new_tokens": 50,
                "repetition_penalty": 1.1
            }
        )
        
        # Create prompt template
        prompt_template = """You are an intelligent document assistant. Use the following context from uploaded documents to answer the question. 
If the answer cannot be found in the context, say "I cannot find this information in the provided documents."
Provide clear and detailed answers based on the context.

Context Information:
{context}

Question: {question}

Please provide a detailed answer based on the context:"""
        
        from langchain.prompts import PromptTemplate
        PROMPT = PromptTemplate(
            template=prompt_template,
            input_variables=["context", "question"]
        )
        
        # Setup RAG chain
        from langchain.chains import RetrievalQA
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=vectorstore.as_retriever(search_kwargs={"k": 3}),
            chain_type_kwargs={"prompt": PROMPT},
            return_source_documents=True
        )
        
        start_time = time.time()
        result = qa_chain.invoke({"query": question})
        response_time = time.time() - start_time
        
        # Format sources
        sources = []
        if "source_documents" in result:
            for i, doc in enumerate(result["source_documents"], 1):
                sources.append({
                    'id': i,
                    'source': doc.metadata.get('source', 'Unknown Document'),
                    'content': doc.page_content[:300] + "..." if len(doc.page_content) > 300 else doc.page_content,
                    'relevance': round((1 - (i / len(result["source_documents"]))) * 0.8 + 0.2, 2)
                })
        
        # Clear memory
        import gc
        gc.collect()
        
        return jsonify({
            'status': 'success',
            'data': {
                'answer': result["result"],
                'sources': sources,
                'metrics': {
                    'response_time': round(response_time, 2),
                    'sources_used': len(sources),
                    'confidence': min(len(sources) / 3 * 100, 100) if sources else 0
                }
            }
        })
        
    except Exception as e:
        print(f"QA error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    print(f"Starting server on port {port}")
    
    # For Render, we need to bind to the port properly
    app.run(host="0.0.0.0", port=port, debug=False)